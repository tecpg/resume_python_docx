from docx import Document
from docx.shared import Inches
import pyttsx3
import qrcode


def speak(text):
    pyttsx3.speak(text)


document = Document()

#User Info
name = input("What is your name?") #"Okeke Chisom Paschal"
speak('Hello ' + name + ' How are you today')
phone = input("What is your phone number?") #"+(234-811-567-77-08)"

email = input("What is your email address?") #"paschalgreen007@gmail.com"

document.add_picture("me.jpg", width=Inches(1.0))
document.add_paragraph(
    name + " | " + phone + " | " + email
)

#About Me
document.add_heading("About me")
document.add_paragraph(input("Tell me about yurself?"))

#Work Experience
document.add_heading("Work Experience")
p = document.add_paragraph()

company = input("Enter Company Name")
from_date = input("From date")
till_present = input("Still Working at "+ company +"?")
end_date = input("End date ")

p.add_run(company + " ").bold = True
p.add_run(from_date + "-" + end_date + '\n').italic = True 

experience_details = input("Describe your role at" + company)
p.add_run(experience_details)

#More Work Experience
while True:
    has_more_experience = input("Do you have more Experiences? (Yes/No)")
    if has_more_experience.lower() == "yes":
            p = document.add_paragraph()

            company = input("Enter Company Name ")
            from_date = input("From date ")
            end_date = input("End date ")

            p.add_run(company + " ").bold = True
            p.add_run(from_date + "-" + end_date + '\n').italic = True 

            experience_details = input("Describe your role at " + company)
            p.add_run(experience_details)
    else:
        break

#Skills
document.add_heading("Skills")
skill = input("Enter Skill")
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    has_more_skills = input("Do you have more Skills? (Yes/No)")
    if has_more_skills.lower() == "yes":
            skill = input("Enter skill ")
            p = document.add_paragraph(skill)
            p.style = 'List Bullet'

    else:
        break

#Tabular data coming soon

# records = (
#     (1, '', ''),
#     (2, '', ''),
#     (3, '', '')
# )

# table = document.add_table(rows=1, cols=3)
# hdr_cells = table.rows[0].cells
# hdr_cells[0].text = 'Experience'
# hdr_cells[1].text = 'Role'
# hdr_cells[2].text = 'Date'
# for qty, id, desc in records:
#     row_cells = table.add_row().cells
#     row_cells[0].text = str(qty)
#     row_cells[1].text = id
#     row_cells[2].text = desc

# document.add_page_break()

#Generate the QR Code
qr_code_filename = "qrcode.jpg"
qrcode_image = qrcode.make('https://github.com/tecpg/resume_python_docx.git')
qrcode_image.save(qr_code_filename)
document.add_picture(qr_code_filename, width=Inches(1.0))
#Footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV written with python docx"

document.save('cv.docx')